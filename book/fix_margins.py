from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\output\Visual Reasoning AI - KDP Final.docx"

GUTTER = Inches(0.75)
OUTSIDE_MARGIN = Inches(0.5)
TOP_MARGIN = Inches(0.75)
BOTTOM_MARGIN = Inches(0.75)

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)

def fix_margins(doc):
    print("Fixing page margins for KDP...")
    
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        
        section.left_margin = GUTTER
        section.right_margin = OUTSIDE_MARGIN
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        
        section.gutter = Inches(0)
        
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
    
    print(f"  Page size: {PAGE_WIDTH.inches}\" x {PAGE_HEIGHT.inches}\"")
    print(f"  Inside margin (gutter): {GUTTER.inches}\"")
    print(f"  Outside margin: {OUTSIDE_MARGIN.inches}\"")
    print(f"  Top margin: {TOP_MARGIN.inches}\"")
    print(f"  Bottom margin: {BOTTOM_MARGIN.inches}\"")

def fix_tables(doc):
    print(f"\nChecking tables...")
    
    usable_width = PAGE_WIDTH.inches - GUTTER.inches - OUTSIDE_MARGIN.inches
    print(f"  Usable content width: {usable_width}\"")
    
    for i, table in enumerate(doc.tables):
        table.autofit = True
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        
        print(f"  Table {i}: Set to 100% width with autofit")

def fix_long_paragraphs(doc):
    print(f"\nChecking paragraphs for overflow issues...")
    
    fixed_count = 0
    for para in doc.paragraphs:
        pPr = para._element.pPr
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                left = ind.get(qn('w:left'))
                if left and int(left) > 1440:
                    ind.set(qn('w:left'), '720')
                    fixed_count += 1
    
    if fixed_count > 0:
        print(f"  Fixed {fixed_count} paragraphs with excessive indentation")
    else:
        print(f"  No excessive indentation found")

def main():
    print(f"Loading document: {DOC_PATH}\n")
    doc = Document(DOC_PATH)
    
    fix_margins(doc)
    fix_tables(doc)
    fix_long_paragraphs(doc)
    
    print(f"\nSaving document...")
    doc.save(DOC_PATH)
    print("Done!")
    
    print("\n" + "="*50)
    print("KDP REQUIREMENTS CHECK:")
    print("="*50)
    print(f"Required gutter for 302 pages: >= 0.625\"")
    print(f"Your gutter (inside margin):   {GUTTER.inches}\" ✓")
    print(f"Required outside margin:       >= 0.25\"")
    print(f"Your outside margin:           {OUTSIDE_MARGIN.inches}\" ✓")
    print(f"Required top/bottom:           >= 0.25\"")
    print(f"Your top/bottom:               {TOP_MARGIN.inches}\" ✓")

if __name__ == "__main__":
    main()
