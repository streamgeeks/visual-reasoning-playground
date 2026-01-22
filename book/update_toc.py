from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\output\Visual Reasoning AI - KDP Final.docx"

CHAPTERS = [
    ("", "Acknowledgments", "i"),
    ("1", "Welcome to Visual Reasoning", ""),
    ("2", "Your First Visual Query", ""),
    ("3", "Drawing Detection Boxes", ""),
    ("4", "Visual Reasoning vs Everything Else", ""),
    ("5", "Models, APIs, and Getting Access", ""),
    ("6", "Your Development Environment", ""),
    ("7", "Auto-Track Any Object", ""),
    ("8", "Smart Counter", ""),
    ("9", "Scene Analyzer", ""),
    ("10", "Zone Monitor", ""),
    ("11", "AI Color Correction Assistant", ""),
    ("12", "Audio Fundamentals", ""),
    ("13", "Intent Extraction", ""),
    ("14", "Multimodal Fusion", ""),
    ("15", "vMix Integration", ""),
    ("16", "OBS Integration", ""),
    ("17", "PTZOptics Advanced", ""),
    ("18", "What Is a Harness", ""),
    ("19", "Agentic Coding", ""),
    ("20", "Logging and Debugging", ""),
    ("21", "Model Swapping", ""),
    ("22", "Sports Broadcasting", ""),
    ("23", "Worship", ""),
    ("24", "Education", ""),
    ("25", "Corporate", ""),
    ("26", "When to Use AI", ""),
    ("27", "Ethics and Privacy", ""),
    ("28", "The Future", ""),
    ("A", "Appendix: Playground Reference", ""),
    ("B", "Appendix: API Quick Reference", ""),
    ("C", "Appendix: Troubleshooting Guide", ""),
    ("D", "Appendix: Glossary", ""),
]

def main():
    print("Loading document...")
    doc = Document(DOC_PATH)
    
    if len(doc.tables) == 0:
        print("No tables found!")
        return
    
    table = doc.tables[0]
    print(f"TOC table has {len(table.rows)} rows")
    
    current_rows = len(table.rows)
    needed_rows = len(CHAPTERS)
    
    if needed_rows > current_rows:
        print(f"Adding {needed_rows - current_rows} rows to table...")
        for _ in range(needed_rows - current_rows):
            table.add_row()
    
    for i, (num, title, page) in enumerate(CHAPTERS):
        if i < len(table.rows):
            row = table.rows[i]
            
            row.cells[0].text = num
            row.cells[1].text = title
            if len(row.cells) > 2:
                row.cells[2].text = page
            
            print(f"  Row {i}: {num} | {title}")
    
    print(f"Saving...")
    doc.save(DOC_PATH)
    print("Done! Table of Contents updated.")

if __name__ == "__main__":
    main()
