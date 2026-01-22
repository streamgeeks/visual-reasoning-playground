from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

OUTPUT_PATH = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\output\Visual Reasoning AI - KDP Final.docx"
CHAPTERS_DIR = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\chapters"

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
INSIDE_MARGIN = Inches(0.875)
OUTSIDE_MARGIN = Inches(0.625)
TOP_MARGIN = Inches(0.75)
BOTTOM_MARGIN = Inches(0.75)

CHAPTER_FILES = [
    "00-front-matter.md",
    "01-welcome-to-visual-reasoning.md",
    "02-your-first-visual-query.md",
    "03-drawing-detection-boxes.md",
    "04-visual-reasoning-vs-everything-else.md",
    "05-models-apis-and-getting-access.md",
    "06-your-development-environment.md",
    "07-auto-track-any-object.md",
    "08-smart-counter.md",
    "09-scene-analyzer.md",
    "10-zone-monitor.md",
    "11-ai-color-correction.md",
    "12-audio-fundamentals.md",
    "13-intent-extraction.md",
    "14-multimodal-fusion.md",
    "15-obs-integration.md",
    "16-ptzoptics-advanced.md",
    "17-vmix-integration.md",
    "18-what-is-a-harness.md",
    "19-agentic-coding.md",
    "20-applied-ideas-proav.md",
    "21-future.md",
    "appendix-a-playground-reference.md",
    "appendix-b-api-reference.md",
    "appendix-c-troubleshooting.md",
    "appendix-d-glossary.md",
]

CHAPTER_TITLES = [
    ("", "Front Matter"),
    ("1", "Welcome to Visual Reasoning"),
    ("2", "Your First Visual Query"),
    ("3", "Drawing Detection Boxes"),
    ("4", "Visual Reasoning vs Everything Else"),
    ("5", "Models, APIs, and Getting Access"),
    ("6", "Your Development Environment"),
    ("7", "Auto-Track Any Object"),
    ("8", "Smart Counter"),
    ("9", "Scene Analyzer"),
    ("10", "Zone Monitor"),
    ("11", "AI Color Correction Assistant"),
    ("12", "Audio Fundamentals"),
    ("13", "Intent Extraction"),
    ("14", "Multimodal Fusion"),
    ("15", "OBS Integration"),
    ("16", "PTZOptics Advanced"),
    ("17", "vMix Integration"),
    ("18", "What Is a Harness"),
    ("19", "Agentic Coding"),
    ("20", "Applied Ideas for ProAV"),
    ("21", "The Future"),
    ("A", "Appendix: Playground Reference"),
    ("B", "Appendix: API Quick Reference"),
    ("C", "Appendix: Troubleshooting Guide"),
    ("D", "Appendix: Glossary"),
]

FONT_NAME = "Garamond"
FONT_SIZE = Pt(12)
HEADING1_SIZE = Pt(18)
HEADING2_SIZE = Pt(14)
HEADING3_SIZE = Pt(12)

def break_long_strings(text):
    def add_breaks(match):
        url = match.group(0)
        broken = url.replace('/', '/\u200b').replace('-', '-\u200b').replace('.', '.\u200b').replace('_', '_\u200b')
        return broken
    
    text = re.sub(r'https?://[^\s]+', add_breaks, text)
    text = re.sub(r'www\.[^\s]+', add_breaks, text)
    text = re.sub(r'github\.com/[^\s]+', add_breaks, text)
    text = re.sub(r'console\.[^\s]+', add_breaks, text)
    text = re.sub(r'api\.[^\s]+', add_breaks, text)
    
    words = text.split()
    result = []
    for word in words:
        if len(word) > 35 and '\u200b' not in word:
            broken = ''
            for i, char in enumerate(word):
                broken += char
                if i > 0 and i % 30 == 0:
                    broken += '\u200b'
            result.append(broken)
        else:
            result.append(word)
    return ' '.join(result)

def clean_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = break_long_strings(text)
    return text

def set_run_font(run, size=None, bold=False):
    run.font.name = FONT_NAME
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    run.font.size = size if size else FONT_SIZE
    run.font.bold = bold

def setup_section(section):
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = INSIDE_MARGIN
    section.right_margin = OUTSIDE_MARGIN
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.gutter = Inches(0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

def add_centered_text(doc, text, size, bold=False, space_after=Pt(12)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size, bold)
    p.paragraph_format.space_after = space_after
    return p

def add_body_text(doc, text, space_after=Pt(6)):
    p = doc.add_paragraph()
    text = break_long_strings(text)
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, HEADING1_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        set_run_font(run, HEADING2_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    else:
        set_run_font(run, HEADING3_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def create_front_matter(doc):
    add_centered_text(doc, "", Pt(36))
    add_centered_text(doc, "", Pt(36))
    add_centered_text(doc, "Visual Reasoning AI", Pt(24), bold=True, space_after=Pt(6))
    add_centered_text(doc, "for Broadcast and ProAV", Pt(18), space_after=Pt(24))
    add_centered_text(doc, "Practical AI Automation for", Pt(14), space_after=Pt(3))
    add_centered_text(doc, "Streaming, ProAV, and Live Production", Pt(14), space_after=Pt(48))
    add_centered_text(doc, "PAUL RICHARDS", Pt(14), bold=True)
    
    add_page_break(doc)
    
    add_centered_text(doc, "", Pt(12))
    add_centered_text(doc, "Copyright 2026 Paul Richards", Pt(11))
    add_centered_text(doc, "All rights reserved.", Pt(11))
    add_centered_text(doc, "", Pt(12))
    add_centered_text(doc, "ISBN: 9798266790568", Pt(11))
    
    add_page_break(doc)
    
    add_centered_text(doc, "DEDICATION", Pt(14), bold=True, space_after=Pt(24))
    add_centered_text(doc, "Dedicated to those who are embracing change.", Pt(12), space_after=Pt(6))
    add_centered_text(doc, "The optimistic pioneers, unwilling to let change", Pt(12), space_after=Pt(6))
    add_centered_text(doc, "get in the way of new innovations.", Pt(12))
    
    add_page_break(doc)
    
    add_centered_text(doc, "CONTENTS", Pt(14), bold=True, space_after=Pt(24))
    
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_p.add_run("Acknowledgments")
    set_run_font(run)
    toc_p.paragraph_format.space_after = Pt(3)
    
    for num, title in CHAPTER_TITLES:
        if num == "":  # Skip front matter in TOC
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        set_run_font(run)
        p.paragraph_format.space_after = Pt(3)
    
    add_page_break(doc)
    
    add_centered_text(doc, "ACKNOWLEDGMENTS", Pt(14), bold=True, space_after=Pt(24))
    add_body_text(doc, "I would like to acknowledge Matthew Davis, Chief Product Engineer at PTZOptics who created the Node-based video tool used in this book.")
    
    add_page_break(doc)

def process_chapter(doc, filepath, is_first_chapter=False):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    paragraph_buffer = []
    in_list = False
    
    def flush_buffer():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            combined = ' '.join(paragraph_buffer)
            combined = clean_markdown(combined)
            if combined.strip():
                add_body_text(doc, combined)
            paragraph_buffer = []
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# Chapter') or line.startswith('# Appendix'):
            flush_buffer()
            if not is_first_chapter:
                add_page_break(doc)
            title = line[2:].strip()
            if 'Appendix' in title:
                add_heading(doc, title.upper(), level=1)
            else:
                title_text = title.replace('Chapter ', '').replace(': ', ' ')
                match = re.match(r'(\d+)[:\s]+(.+)', title_text)
                if match:
                    num = match.group(1)
                    ttl = match.group(2).upper()
                    add_heading(doc, f"{num} {ttl}", level=1)
                else:
                    add_heading(doc, title.upper(), level=1)
            in_list = False
            
        elif line.startswith('## '):
            flush_buffer()
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            add_heading(doc, line[3:].strip(), level=2)
            in_list = False
            
        elif line.startswith('### '):
            flush_buffer()
            add_heading(doc, line[4:].strip(), level=3)
            in_list = False
            
        elif line.startswith('- '):
            flush_buffer()
            text = clean_markdown(line[2:].strip())
            text = break_long_strings(text)
            p = doc.add_paragraph()
            run = p.add_run(f"  \u2022  {text}")
            set_run_font(run)
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(3)
            in_list = True
            
        elif re.match(r'^\d+\.\s', line):
            flush_buffer()
            match = re.match(r'^(\d+)\.\s(.+)', line)
            if match:
                num = match.group(1)
                text = clean_markdown(match.group(2).strip())
                p = doc.add_paragraph()
                run = p.add_run(f"  {num}. {text}")
                set_run_font(run)
                p.paragraph_format.left_indent = Inches(0.15)
                p.paragraph_format.space_after = Pt(3)
            in_list = True
            
        elif line.startswith('|'):
            flush_buffer()
            continue
            
        elif line.startswith('---'):
            flush_buffer()
            in_list = False
            
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            flush_buffer()
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            in_list = False
            
        elif line.strip():
            if in_list:
                flush_buffer()
                in_list = False
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
            
            text = line.strip()
            if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
                flush_buffer()
                p = doc.add_paragraph()
                run = p.add_run(clean_markdown(text))
                set_run_font(run)
                run.italic = True
                p.paragraph_format.space_after = Pt(6)
            else:
                paragraph_buffer.append(text)
        else:
            if paragraph_buffer:
                flush_buffer()
                if in_list:
                    in_list = False
    
    flush_buffer()

def main():
    print("Creating new KDP-compliant document from scratch...")
    doc = Document()
    
    print("Setting up page margins...")
    setup_section(doc.sections[0])
    
    print("Creating front matter...")
    create_front_matter(doc)
    
    print("Adding chapters...")
    for idx, chapter_file in enumerate(CHAPTER_FILES):
        filepath = os.path.join(CHAPTERS_DIR, chapter_file)
        if not os.path.exists(filepath):
            print(f"  WARNING: {chapter_file} not found")
            continue
        print(f"  {chapter_file}")
        process_chapter(doc, filepath, is_first_chapter=(idx == 0))
    
    print("Verifying all sections have correct margins...")
    for section in doc.sections:
        setup_section(section)
    
    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    
    print("\n" + "="*55)
    print("KDP MARGIN REQUIREMENTS (305 pages, 6x9 book):")
    print("="*55)
    print(f"  Inside (gutter):  Required >= 0.625\"  |  Set: {INSIDE_MARGIN.inches}\"")
    print(f"  Outside margin:   Required >= 0.25\"   |  Set: {OUTSIDE_MARGIN.inches}\"")  
    print(f"  Top margin:       Required >= 0.25\"   |  Set: {TOP_MARGIN.inches}\"")
    print(f"  Bottom margin:    Required >= 0.25\"   |  Set: {BOTTOM_MARGIN.inches}\"")
    print("="*55)
    print("Done! Document created from scratch with safe margins.")

if __name__ == "__main__":
    main()
