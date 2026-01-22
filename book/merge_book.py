from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

TEMPLATE_PATH = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\output\Visual Reasoning_  Use this format.docx"
OUTPUT_PATH = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\output\Visual Reasoning AI - KDP Final.docx"
CHAPTERS_DIR = r"C:\Users\paulw\OneDrive\Desktop\VisualReasoning\book\chapters"

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
INSIDE_MARGIN = Inches(0.75)
OUTSIDE_MARGIN = Inches(0.5)
TOP_MARGIN = Inches(0.75)
BOTTOM_MARGIN = Inches(0.75)

CHAPTER_FILES = [
    "01-welcome-to-visual-reasoning.md",
    "02-your-first-visual-query.md",
    "03-drawing-detection-boxes.md",
    "04-visual-reasoning-vs-everything-else.md",
    "05-models-apis-and-getting-access.md",
    "06-your-development-environment.md",
    "07-auto-track-any-object.md",
    "08-smart-counter.md",
    "09-scene-analyzer.md",
    "10-zone-monitor.md",
    "11-ai-color-correction.md",
    "12-audio-fundamentals.md",
    "13-intent-extraction.md",
    "14-multimodal-fusion.md",
    "15-vmix-integration.md",
    "16-obs-integration.md",
    "17-ptzoptics-advanced.md",
    "18-what-is-a-harness.md",
    "19-agentic-coding.md",
    "20-logging-debugging.md",
    "21-model-swapping.md",
    "22-sports-broadcasting.md",
    "23-worship.md",
    "24-education.md",
    "25-corporate.md",
    "26-when-to-use-ai.md",
    "27-ethics-privacy.md",
    "28-future.md",
    "appendix-a-playground-reference.md",
    "appendix-b-api-reference.md",
    "appendix-c-troubleshooting.md",
    "appendix-d-glossary.md",
]

FONT_NAME = "Garamond"
FONT_SIZE = Pt(12)
HEADING1_SIZE = Pt(18)
HEADING2_SIZE = Pt(14)
HEADING3_SIZE = Pt(12)

def clean_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def set_run_font(run, size=None, bold=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = size if size else FONT_SIZE
    run.font.bold = bold

def set_page_margins(doc):
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.left_margin = INSIDE_MARGIN
        section.right_margin = OUTSIDE_MARGIN
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.gutter = Inches(0)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

def fix_all_tables(doc):
    for table in doc.tables:
        table.autofit = True
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(OxmlElement('w:br'))
    run._element[-1].set(qn('w:type'), 'page')

def add_paragraph_with_font(doc, text, size=None, bold=False, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, bold)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    
    if level == 1:
        set_run_font(run, HEADING1_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        set_run_font(run, HEADING2_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    else:
        set_run_font(run, HEADING3_SIZE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    return p

def create_list_styles(doc):
    try:
        styles = doc.styles
        style_names = [s.name for s in styles]
        if 'List Bullet' not in style_names:
            bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = FONT_NAME
            bullet_style.font.size = FONT_SIZE
        if 'List Number' not in style_names:
            num_style = styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
            num_style.font.name = FONT_NAME
            num_style.font.size = FONT_SIZE
    except:
        pass

def process_chapter(doc, filepath, is_first_chapter=False):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    paragraph_buffer = []
    in_list = False
    
    def flush_paragraph_buffer():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            combined = ' '.join(paragraph_buffer)
            combined = clean_markdown(combined)
            if combined.strip():
                add_paragraph_with_font(doc, combined)
            paragraph_buffer = []
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# Chapter') or line.startswith('# Appendix'):
            flush_paragraph_buffer()
            if not is_first_chapter:
                add_page_break(doc)
            title = line[2:].strip()
            if 'Appendix' in title:
                add_heading(doc, title.upper(), level=1)
            else:
                title_text = title.replace('Chapter ', '').replace(': ', ' ')
                match = re.match(r'(\d+)[:\s]+(.+)', title_text)
                if match:
                    num = match.group(1)
                    ttl = match.group(2).upper()
                    add_heading(doc, f"{num} {ttl}", level=1)
                else:
                    add_heading(doc, title.upper(), level=1)
            in_list = False
            
        elif line.startswith('## '):
            flush_paragraph_buffer()
            doc.add_paragraph()
            title = line[3:].strip()
            add_heading(doc, title, level=2)
            in_list = False
            
        elif line.startswith('### '):
            flush_paragraph_buffer()
            title = line[4:].strip()
            add_heading(doc, title, level=3)
            in_list = False
            
        elif line.startswith('- '):
            flush_paragraph_buffer()
            text = clean_markdown(line[2:].strip())
            p = doc.add_paragraph()
            run = p.add_run(f"\u2022  {text}")
            set_run_font(run)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.right_indent = Inches(0)
            p.paragraph_format.space_after = Pt(3)
            in_list = True
            
        elif re.match(r'^\d+\.\s', line):
            flush_paragraph_buffer()
            match = re.match(r'^(\d+)\.\s(.+)', line)
            if match:
                num = match.group(1)
                text = clean_markdown(match.group(2).strip())
                p = doc.add_paragraph()
                run = p.add_run(f"{num}.  {text}")
                set_run_font(run)
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.first_line_indent = Inches(-0.15)
                p.paragraph_format.right_indent = Inches(0)
                p.paragraph_format.space_after = Pt(3)
            in_list = True
            
        elif line.startswith('|'):
            flush_paragraph_buffer()
            continue
            
        elif line.startswith('---'):
            flush_paragraph_buffer()
            in_list = False
            
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            flush_paragraph_buffer()
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.right_indent = Inches(0)
            in_list = False
            
        elif line.strip():
            if in_list:
                flush_paragraph_buffer()
                in_list = False
                doc.add_paragraph()
            
            text = line.strip()
            if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
                flush_paragraph_buffer()
                p = doc.add_paragraph()
                run = p.add_run(clean_markdown(text))
                set_run_font(run)
                run.italic = True
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.right_indent = Inches(0)
            else:
                paragraph_buffer.append(text)
        else:
            if paragraph_buffer:
                flush_paragraph_buffer()
                if in_list:
                    in_list = False
                    doc.add_paragraph()
    
    flush_paragraph_buffer()

def main():
    print("Loading template document...")
    doc = Document(TEMPLATE_PATH)
    
    print("Setting KDP-compliant margins (6x9 book, 302+ pages)...")
    set_page_margins(doc)
    
    create_list_styles(doc)
    
    start_index = None
    in_acknowledgments = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "ACKNOWLEDGMENTS":
            in_acknowledgments = True
            continue
        if in_acknowledgments and re.match(r'^\d+\s+[A-Z]', text):
            start_index = i
            print(f"Found first chapter at paragraph {i}: {text[:50]}...")
            break
    
    if start_index is None:
        print("ERROR: Could not find where chapters start in template")
        return
    
    print(f"Removing placeholder chapters...")
    paragraphs_to_delete = list(range(start_index, len(doc.paragraphs)))
    for i in reversed(paragraphs_to_delete):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    print(f"Removed {len(paragraphs_to_delete)} placeholder paragraphs")
    
    print("Adding chapter content...")
    for idx, chapter_file in enumerate(CHAPTER_FILES):
        filepath = os.path.join(CHAPTERS_DIR, chapter_file)
        if not os.path.exists(filepath):
            print(f"WARNING: {chapter_file} not found")
            continue
        print(f"  {chapter_file}")
        process_chapter(doc, filepath, is_first_chapter=(idx == 0))
    
    print("Fixing table widths...")
    fix_all_tables(doc)
    
    print("Final margin check...")
    set_page_margins(doc)
    
    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    
    print("\n" + "="*50)
    print("KDP REQUIREMENTS FOR 302-PAGE BOOK:")
    print("="*50)
    print(f"  Gutter (inside):  Required >= 0.625\"  |  Set: {INSIDE_MARGIN.inches}\"")
    print(f"  Outside margin:   Required >= 0.25\"   |  Set: {OUTSIDE_MARGIN.inches}\"")
    print(f"  Top/Bottom:       Required >= 0.25\"   |  Set: {TOP_MARGIN.inches}\"")
    print(f"  Page size:        6\" x 9\"")
    print("="*50)
    print("Done!")

if __name__ == "__main__":
    main()
